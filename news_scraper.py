import feedparser
import time
import argparse
import docx
import requests
from io import BytesIO
import datetime
import os
from bs4 import BeautifulSoup
from dateutil import parser as date_parser
from dateutil.tz import tzlocal

def clean_html(raw_html):
    if not raw_html:
        return "No summary available."
    soup = BeautifulSoup(raw_html, "html.parser")
    return soup.get_text(separator=' ', strip=True)

def generate_html_card(color, title, link, author, published, image_url, description, paywall, stance):
    img_tag = f'<div class="h-48 w-full border-b-2 border-{color} overflow-hidden relative"><img src="{image_url}" alt="" class="w-full h-full object-cover grayscale mix-blend-multiply hover:mix-blend-normal hover:grayscale-0 transition-all duration-500"></div>' if image_url else f'<div class="h-48 w-full border-b-2 border-{color} bg-surface-container-highest flex items-center justify-center font-headline text-{color}/50 font-black text-2xl relative">NO IMAGE</div>'
    
    author_str = f"BY {author}" if author else "UNKNOWN"
    
    paywall_color = "bg-[#ff6e84] text-white" if "PAYWALL" in paywall else "bg-[#cafd00] text-black"
    
    return f"""
    <div class="bg-surface-container-high border-2 border-{color} poster-shadow hover:-translate-y-2 hover:-translate-x-2 transition-all group flex flex-col h-full relative">
        <div class="absolute top-2 right-2 z-10 flex flex-col gap-1 items-end">
            <span class="px-2 py-0.5 text-[9px] font-headline font-black uppercase tracking-widest border border-{color} bg-black/80 text-{color} backdrop-blur-sm backdrop-saturate-150">{stance}</span>
            <span class="px-2 py-0.5 text-[9px] font-headline font-black uppercase tracking-widest {paywall_color}">{paywall}</span>
        </div>
        {img_tag}
        <div class="p-6 flex-1 flex flex-col">
            <h3 class="text-xl font-black font-headline text-white uppercase mb-4 group-hover:text-{color} transition-colors line-clamp-3">{title}</h3>
            <p class="text-sm font-body text-white/70 mb-6 flex-1 line-clamp-4">{description}</p>
            <div class="mt-auto border-t-2 border-{color}/30 pt-4 flex flex-wrap gap-y-2 justify-between items-center">
                <div class="flex flex-col">
                    <span class="text-xs font-headline font-bold text-{color} uppercase">{published[:16]}</span>
                    <span class="text-[10px] font-headline text-white/50 uppercase">{author_str}</span>
                </div>
                <a href="{link}" target="_blank" class="bg-{color} text-black px-4 py-2 font-headline font-black text-sm uppercase skew-x-[-10deg] hover:scale-110 shadow-[4px_4px_0px_0px_rgba(0,0,0,1)] transition-transform">READ</a>
            </div>
        </div>
    </div>
    """

def is_today(date_string):
    if not date_string:
        return False
    try:
        dt = date_parser.parse(date_string)
        if dt.tzinfo:
            dt = dt.astimezone(tzlocal())
        return dt.date() == datetime.datetime.now().date()
    except Exception:
        today_str = datetime.datetime.now().strftime("%d %b %Y")
        return today_str in date_string

def fetch_and_append_news(source, document, fetch_all_today=True, max_items=5):
    feed_name = source['name']
    feed_url = source['url']
    color = source['color']
    paywall = source.get('paywall', 'UNKNOWN')
    stance = source.get('stance', 'UNKNOWN')
    
    print(f"\n{'='*50}")
    print(f"Fetching News from: {feed_name}")
    print(f"{'='*50}\n")
    
    document.add_heading(f"Source: {feed_name} | {stance} | {paywall}", level=1)
    
    html_cards = []
    events = []
    
    try:
        feed = feedparser.parse(feed_url)
        if feed.bozo and not feed.entries:
            print(f"Error parsing feed for {feed_name}. Skipping.")
            return html_cards, events

        entries = feed.entries
        valid_entries = []

        if fetch_all_today:
            for entry in entries:
                published = entry.get('published', '')
                if is_today(published):
                    valid_entries.append(entry)
        else:
            valid_entries = entries[:max_items]
            
        if not valid_entries:
            print("No new news items found for today.")
            return html_cards, events

        print(f"-> Found {len(valid_entries)} articles.")

        for i, entry in enumerate(valid_entries, 1):
            title = entry.get('title', 'No Title')
            link = entry.get('link', 'No Link')
            author = entry.get('author', entry.get('dc_creator', 'Unknown Writer'))
            raw_summary = entry.get('description', entry.get('summary', ''))
            description = clean_html(raw_summary)
            published = entry.get('published', 'No publish date')
            
            image_url = None
            if 'media_content' in entry and len(entry.media_content) > 0:
                image_url = entry.media_content[0].get('url')
            elif 'media_thumbnail' in entry and len(entry.media_thumbnail) > 0:
                image_url = entry.media_thumbnail[0].get('url')
            elif 'links' in entry:
                for lnk in entry.links:
                    if lnk.get('type', '').startswith('image/') and 'href' in lnk:
                        image_url = lnk['href']
                        break

            print(f"[{i}] {title}")
            
            try:
                dt = date_parser.parse(published)
                if dt.tzinfo:
                    dt = dt.astimezone(tzlocal())
                event_dt = dt
            except Exception:
                event_dt = datetime.datetime.now()

            events.append({
                'title': title,
                'link': link,
                'description': description,
                'published_str': published[:16],
                'dt': event_dt,
                'source_name': feed_name,
                'source_short': source['short'],
                'color': color,
                'paywall': paywall,
                'stance': stance
            })
            
            # --- WORD DOCUMENT GENERATION ---
            document.add_heading(title, level=2)
            if author:
                p_auth = document.add_paragraph()
                p_auth.add_run(f"By {author} | {published}").italic = True
                
            if image_url:
                try:
                    response = requests.get(image_url, timeout=5)
                    if response.status_code == 200:
                        image_stream = BytesIO(response.content)
                        document.add_picture(image_stream, width=docx.shared.Inches(4))
                except Exception:
                    pass

            document.add_paragraph(description)
            p_link = document.add_paragraph("Read more: ")
            p_link.add_run(link).underline = True
            document.add_paragraph("-" * 50)
            
            # --- HTML GENERATION ---
            html_cards.append(generate_html_card(color, title, link, author, published, image_url, description, paywall, stance))

    except Exception as e:
        print(f"Failed to fetch news from {feed_name}: {e}")
        
    return html_cards, events

def build_timeline_html(all_events):
    import re
    # Removed standalone 'us', 'usa', 'u.s.', 'united states' because they trigger on domestic news
    keywords = ['israel', 'iran', 'lebanon', 'palestine', 'gaza', 'hamas', 'hezbollah', 'yemen', 'houthi', 'syria', 'iraq', 'tehran', 'tel aviv', 'jerusalem', 'idf', 'middle east', 'jordan', 'egypt', 'saudi', 'qatar', 'uae', 'gulf']
    
    filtered_events = []
    for e in all_events:
        text_to_search = (e['title'] + " " + e['description']).lower()
        found = False
        for kw in keywords:
            if kw in ['idf', 'uae', 'gulf']:
                if re.search(r'\b' + re.escape(kw) + r'\b', text_to_search):
                    found = True; break
            else:
                if kw in text_to_search:
                    found = True; break
        if found:
            filtered_events.append(e)
            
    if not filtered_events:
        return ""
        
    filtered_events.sort(key=lambda x: x['dt'])
    
    nodes_html = ""
    for idx, e in enumerate(filtered_events):
        color = e['color']
        justify = "justify-start" if idx % 2 == 0 else "justify-end"
        text_align = "text-left pr-8 md:pr-16" if idx % 2 == 0 else "text-right pl-8 md:pl-16"
        align_items = "items-start" if idx % 2 == 0 else "items-end"
        
        dot = f'<div class="absolute left-1/2 top-1/2 -translate-x-1/2 -translate-y-1/2 w-4 h-4 md:w-6 md:h-6 bg-black border-4 border-{color} rotate-45 z-20 shadow-[0_0_15px_var(--tw-shadow-color)] shadow-{color} group-hover:bg-{color} transition-colors"></div>'
        
        paywall_color = "bg-[#ff6e84] text-white" if "PAYWALL" in e['paywall'] else "bg-[#cafd00] text-black"
        tags = f"""
        <div class="flex gap-2 {justify} mb-3 flex-wrap">
            <span class="px-2 py-0.5 text-[9px] font-headline font-black uppercase tracking-widest border border-{color} bg-black/80 text-{color}">{e['stance']}</span>
            <span class="px-2 py-0.5 text-[9px] font-headline font-black uppercase tracking-widest {paywall_color}">{e['paywall']}</span>
        </div>
        """
        
        nodes_html += f"""
        <div class="relative flex w-full {justify} items-center mb-16 group">
            {dot}
            <div class="w-1/2 flex flex-col {align_items} {text_align} relative z-10 transition-transform duration-300 hover:scale-[1.02] hover:-translate-y-1">
                <div class="inline-block px-3 py-1 bg-{color} text-[black] font-headline font-black text-xs uppercase skew-x-[-10deg] mb-3 shadow-[4px_4px_0px_transparent] group-hover:shadow-[4px_4px_0px_#fff] transition-all">{e['published_str']} | {e['source_short']}</div>
                {tags}
                <h3 class="text-xl md:text-3xl font-black font-headline text-white uppercase mb-3 group-hover:text-{color} transition-colors"><a href="{e['link']}" target="_blank">{e['title']}</a></h3>
                <p class="text-sm md:text-base font-body text-white/60 line-clamp-3 bg-black/50 p-4 border-l-4 border-{color} group-hover:bg-white/5 transition-colors text-left">{e['description']}</p>
            </div>
        </div>
        """

    return f"""
    <!-- WAR TIMELINE -->
    <section id="timeline" class="relative py-24 bg-[#050505] border-b-[16px] border-white overflow-hidden concrete-texture">
        <div class="max-w-7xl mx-auto px-4 md:px-8 relative z-10">
            <div class="text-center mb-20 relative">
                <h2 class="text-6xl md:text-8xl lg:text-[10rem] font-black font-headline text-white uppercase leading-none skew-x-[-10deg] glitch-hover drop-shadow-[8px_8px_0px_rgba(255,255,255,0.1)]">
                    WAR<br/><span class="text-error animate-pulse-fast">TIMELINE</span>
                </h2>
                <div class="inline-block bg-white text-black px-6 py-2 font-headline font-black uppercase text-xl mt-6 skew-x-[-10deg] poster-shadow">MIDDLE EAST CONFLICT</div>
                <p class="mt-6 text-white/50 font-body text-sm max-w-2xl mx-auto uppercase tracking-widest border-t border-white/20 pt-4">Chronological analysis derived from live intercepted broadcasts across Global Networks.</p>
            </div>
            
            <div class="relative w-full py-10 before:content-[''] before:absolute before:left-1/2 before:top-0 before:bottom-0 before:w-1 before:bg-white/20 before:-translate-x-1/2 before:z-0">
                {nodes_html}
            </div>
        </div>
    </section>
    """

def generate_section_html(source, cards_html_list):
    if not cards_html_list:
        return ""
        
    color = source['color']
    short = source['short']
    name = source['name']
    
    cards_str = "".join(cards_html_list)
    
    return f"""
    <!-- SECTION: {short} -->
    <section id="{short.lower()}" class="relative min-h-[400px] flex flex-col justify-center px-8 md:px-16 py-20 bg-surface border-b-[12px] border-{color} overflow-hidden">
        <div class="absolute top-10 right-10 opacity-5 rotate-12 scale-150 pointer-events-none">
            <span class="text-[15rem] md:text-[20rem] font-black text-{color} font-headline whitespace-nowrap">{short}</span>
        </div>
        <div class="relative z-10 w-full max-w-7xl mx-auto mb-16">
            <div class="bg-{color} text-[black] px-6 py-2 font-headline font-black uppercase text-2xl mb-8 inline-block skew-x-[-10deg] poster-shadow hover:-rotate-3 hover:scale-110 transition-transform cursor-pointer animate-pulse-fast glitch-hover">{name}</div>
            <h2 class="text-6xl md:text-8xl font-black font-headline text-{color} leading-[0.9] tracking-tighter uppercase drip-effect glitch-hover">
                LATEST <br/>INTEL
            </h2>
        </div>
        
        <div class="relative z-10 grid grid-cols-1 md:grid-cols-2 lg:grid-cols-3 xl:grid-cols-4 gap-8 w-full max-w-7xl mx-auto">
            {cards_str}
        </div>
    </section>
    """

def get_html_head_and_nav():
    return """<!DOCTYPE html>
<html class="dark" lang="en"><head>
<meta charset="utf-8"/>
<meta content="width=device-width, initial-scale=1.0" name="viewport"/>
<title>URBAN NEWS | Vandalize the Truth</title>
<script src="https://cdn.tailwindcss.com?plugins=forms,container-queries"></script>
<link href="https://fonts.googleapis.com/css2?family=Space+Grotesk:wght@300;400;700;900&amp;family=Epilogue:wght@300;400;700;900&amp;display=swap" rel="stylesheet"/>
<link href="https://fonts.googleapis.com/css2?family=Material+Symbols+Outlined:wght,FILL@100..700,0..1&amp;display=swap" rel="stylesheet"/>
<script id="tailwind-config">
    tailwind.config = {
        darkMode: "class",
        theme: {
            extend: {
                colors: {
                    "surface": "#0e0e0f",
                    "primary": "#f3ffca",
                    "secondary": "#ff6b9b",
                    "tertiary": "#81ecff",
                    "error": "#ff6e84",
                    "primary-fixed": "#cafd00",
                    "surface-container-high": "#1f1f21",
                    "surface-container-highest": "#262627",
                    "background": "#0e0e0f",
                    "on-background": "#ffffff",
                },
                fontFamily: {
                    "headline": ["Space Grotesk"],
                    "body": ["Epilogue"],
                },
                animation: {
                    'glitch': 'glitch 1s linear infinite',
                    'pulse-fast': 'pulse 1.5s cubic-bezier(0.4, 0, 0.6, 1) infinite',
                    'float': 'float 3s ease-in-out infinite',
                    'marquee': 'marquee 25s linear infinite',
                },
                keyframes: {
                    glitch: {
                        '0%, 100%': { transform: 'translate(0)' },
                        '33%': { transform: 'translate(-2px, 2px)', textShadow: '2px 0 #ff6b9b' },
                        '66%': { transform: 'translate(2px, -2px)', textShadow: '-2px 0 #81ecff' },
                    },
                    float: {
                        '0%, 100%': { transform: 'translateY(0)' },
                        '50%': { transform: 'translateY(-10px)' },
                    },
                    marquee: {
                        '0%': { transform: 'translateX(0)' },
                        '100%': { transform: 'translateX(-50%)' },
                    }
                }
            }
        }
    }
</script>
<style>
    .concrete-texture {
        background-image: radial-gradient(circle at 2px 2px, rgba(255,255,255,0.05) 1px, transparent 0);
        background-size: 4px 4px;
    }
    .poster-shadow { box-shadow: 8px 8px 0px 0px rgba(0,0,0,1); }
    .drip-effect::after, .drip::after {
        content: ''; position: absolute; bottom: -20px; left: 10%; width: 8px; height: 40px;
        background: currentColor; border-radius: 0 0 10px 10px;
        box-shadow: 20px 10px 0 currentColor, 50px -5px 0 currentColor, 80px 15px 0 currentColor;
    }
    html { scroll-behavior: smooth; }
    .stencil-text { -webkit-text-stroke: 2px currentColor; color: transparent; }
    .glitch-hover:hover { animation: glitch 0.2s linear infinite; }
    .marquee-container { overflow: hidden; white-space: nowrap; }
    .news-card { transition: all 0.3s cubic-bezier(0.175, 0.885, 0.32, 1.275); }
    .news-card:hover { transform: scale(1.02) rotate(-1deg); box-shadow: 15px 15px 0 0 currentColor; }
    .skew-heavy { transform: skewx(-12deg); }
    .skew-reverse { transform: skewx(12deg); }
</style>
</head>
<body class="bg-background text-on-background font-body concrete-texture">

<!-- MARQUEE TICKER -->
<div class="bg-secondary text-white py-2 overflow-hidden border-b-4 border-black z-[60] sticky top-0 marquee-container">
    <div class="inline-block animate-marquee uppercase font-headline font-black tracking-widest text-sm">
        SYSTEM BREAKDOWN IN PROGRESS • DISCONNECT THE GRID • TRUTH IS VANDALISM • NO PERMISSION REQUIRED • THE WALLS HAVE EYES • BROADCASTING FROM THE UNDERGROUND • 
        SYSTEM BREAKDOWN IN PROGRESS • DISCONNECT THE GRID • TRUTH IS VANDALISM • NO PERMISSION REQUIRED • THE WALLS HAVE EYES • BROADCASTING FROM THE UNDERGROUND • 
    </div>
</div>

<!-- TopNavBar -->
<header class="sticky top-10 z-50 flex justify-between items-center px-6 py-4 w-full bg-[#0e0e0f] border-b-4 border-[#ff6b9b] shadow-[8px_8px_0px_0px_rgba(255,107,155,1)]">
    <div class="text-4xl font-black text-[#f3ffca] tracking-tighter skew-x-[-10deg] font-headline uppercase animate-glitch">URBAN NEWS</div>
    <nav class="hidden md:flex flex-wrap items-center gap-2">
        <a class="font-headline font-bold tracking-tighter uppercase text-[#0e0e0f] bg-white border-2 border-black px-4 py-1 skew-x-[-12deg] hover:scale-110 duration-100 transition-all shadow-[4px_4px_0_0_#ff007f] animate-pulse-fast mr-4" href="#timeline">TIMELINE</a>
        <a class="font-headline font-bold tracking-tighter uppercase text-[#0e0e0f] bg-primary px-4 py-1 skew-x-[-12deg] hover:scale-110 duration-100 transition-all" href="#nyt">NYT</a>
        <a class="font-headline font-bold tracking-tighter uppercase text-[#0e0e0f] bg-secondary px-4 py-1 skew-x-[-12deg] hover:scale-110 duration-100 transition-all" href="#guardian">Guardian</a>
        <a class="font-headline font-bold tracking-tighter uppercase text-[#0e0e0f] bg-tertiary px-4 py-1 skew-x-[-12deg] hover:scale-110 duration-100 transition-all" href="#bbc">BBC</a>
        <a class="font-headline font-bold tracking-tighter uppercase text-[#0e0e0f] bg-error px-4 py-1 skew-x-[-12deg] hover:scale-110 duration-100 transition-all" href="#wapo">WaPo</a>
        <a class="font-headline font-bold tracking-tighter uppercase text-[#0e0e0f] bg-primary-fixed px-4 py-1 skew-x-[-12deg] hover:scale-110 duration-100 transition-all" href="#al jazeera">Al Jazeera</a>
    </nav>
</header>
<main class="w-full min-h-screen">
"""

def get_html_footer():
    return """
</main>
<footer class="flex flex-col md:flex-row justify-between items-center px-10 py-12 w-full bg-[#0e0e0f] border-t-8 border-dotted border-[#81ecff]">
    <div class="mb-8 md:mb-0">
        <div class="text-[#f3ffca] font-black text-3xl font-headline uppercase mb-2 animate-glitch skew-heavy">URBAN NEWS</div>
        <p class="font-body text-xs font-black uppercase tracking-widest text-[#ff6b9b]">© NO PERMISSION REQUIRED</p>
    </div>
    <div class="flex flex-wrap justify-center gap-8">
        <a class="font-body text-xs font-black uppercase tracking-widest text-white/50 hover:text-[#f3ffca] transition-all underline decoration-4" href="#">Automated Scraping</a>
        <a class="font-body text-xs font-black uppercase tracking-widest text-white/50 hover:text-[#f3ffca] transition-all" href="#">Vandalize The Truth</a>
    </div>
</footer>
</body></html>
"""

def main():
    parser = argparse.ArgumentParser(description="News Scraper generating Word and Urban HTML Reports.")
    parser.add_argument("--all-today", action="store_true", default=True, help="Fetch all news from today.")
    parser.add_argument("-n", "--number", type=int, default=5, help="Limit items if all-today is false.")
    args = parser.parse_args()

    # Sources mapped to URBAN theme colors, plus political stance and paywall data
    news_sources = [
        {"name": "The New York Times", "url": "https://rss.nytimes.com/services/xml/rss/nyt/World.xml", "color": "primary", "short": "NYT", "paywall": "METERED PAYWALL", "stance": "CENTER-LEFT"},
        {"name": "The Guardian", "url": "https://www.theguardian.com/world/rss", "color": "secondary", "short": "Guardian", "paywall": "FREE (DONATION)", "stance": "LEFT/PROGRESSIVE"},
        {"name": "BBC News", "url": "http://feeds.bbci.co.uk/news/world/rss.xml", "color": "tertiary", "short": "BBC", "paywall": "FREE", "stance": "CENTER"},
        {"name": "Washington Post", "url": "https://feeds.washingtonpost.com/rss/world", "color": "error", "short": "WaPo", "paywall": "METERED PAYWALL", "stance": "CENTER-LEFT"},
        {"name": "Al Jazeera", "url": "https://www.aljazeera.com/xml/rss/all.xml", "color": "primary-fixed", "short": "Al Jazeera", "paywall": "FREE", "stance": "CENTER-LEFT"}
    ]
    
    base_dir = os.path.dirname(os.path.abspath(__file__))
    reports_dir = os.path.join(base_dir, "reports")
    os.makedirs(reports_dir, exist_ok=True)
    
    today_str = datetime.datetime.now().strftime("%Y-%m-%d")
    docx_out_path = os.path.join(reports_dir, f"urban_news_{today_str}.docx")
    html_out_path = os.path.join(reports_dir, f"urban_news_{today_str}.html")

    doc = docx.Document()
    doc.add_heading(f"Urban News Report - {today_str}", 0)
    
    html_sections = []
    all_events = []
    
    print(f"Starting Scraper with Urban Theme ({today_str})...")
    
    for source in news_sources:
        # Fetch cards for this specific source
        cards, events = fetch_and_append_news(source, doc, fetch_all_today=args.all_today, max_items=args.number)
        all_events.extend(events)
        
        # Build Section HTML
        section_html = generate_section_html(source, cards)
        html_sections.append(section_html)
        time.sleep(1)
        
    timeline_html = build_timeline_html(all_events)
        
    try:
        doc.save(docx_out_path)
        print(f"\nSaved DOCX report to: {docx_out_path}")
    except Exception as e:
        pass

    try:
        final_html = get_html_head_and_nav() + "".join(html_sections) + timeline_html + get_html_footer()
        with open(html_out_path, "w", encoding="utf-8") as f:
            f.write(final_html)
        print(f"Saved URBAN HTML dashboard to: {html_out_path}")
    except Exception as e:
        print(f"Failed to save HTML: {e}")

if __name__ == "__main__":
    main()
